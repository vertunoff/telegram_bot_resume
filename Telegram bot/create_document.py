from docx import Document
from docx.shared import Inches, Pt
def add_run(paragraph, text='', font='Times new Roman', size=14, bold=False, italic=False, underline=False, space=False, newline=False, **kwargs):
    if newline:
        text+='\n'
    if space:
        text+= ' '
    paragraph.add_run(text)
    run = paragraph.runs[-1]
    run.bold = bold
    run.underline = underline
    font = run.font
    font.size = Pt(size)
    font.italic = italic

def newline(paragraph):
    add_run(paragraph, '\n')

def complete_profile(document, tags : list, data : list, photo : str):
    name = data.pop(0)
    row_cells = document.tables[0].rows[0].cells
    paragraph = row_cells[0].paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(photo, height = Inches(0.72*4/1.4))
    
    coll_cells = document.tables[0].columns[1].cells
    profile = coll_cells[0].paragraphs[0]
    add_run(profile, name, size=18, bold=True, newline=True)
    add_run(profile, '\n', size=8)
    for tag, value in zip(tags, data):
        add_run(profile, f'{tag}: {value}')
        if tag != 'Электронная почта':
            add_run(profile, '\n', 22)

def complete_personal(document, tags : list, data : list):
    document.add_paragraph()
    personal = document.paragraphs[-1]
    add_run(personal, '\t• Личная информация', bold=True, size=18, newline=True)
    for tag, value in zip(tags, data):
        add_run(personal, f'{tag}: {value}')
        if tag != 'Семейное положение':
            newline(personal)

def complete_exp(document, tags : list, data : list, exp=False):
    document.add_paragraph()
    experience = document.paragraphs[-1]
    add_run(experience, '\t• Опыт работы', bold=True, size=18, newline=exp)
    if not exp:
        add_run(experience, ': нет', bold=True, size=18)
        return
    for tag, value in zip(tags, data):
        add_run(experience, f'{tag}: {value}')
        if tag != 'Должностные обязанности и достижения':
            newline(experience)

def complete_education(document, tags : list, data : list):
    document.add_paragraph()
    education = document.paragraphs[-1]
    add_run(education, '\t• Образование', bold=True, size=18, newline=bool(data))
    if not data:
        add_run(education, ': нет', bold=True, size=18)
        return
    for tag, value in zip(tags, data):
        add_run(education, f'{tag}: {value}')
        if tag != 'Форма обучения':
            newline(education)

def complete_optional(document, tags : list, data : list):
    document.add_paragraph()
    optional = document.paragraphs[-1]
    add_run(optional, '\t• Дополнительная информация', bold=True, size=18, newline=True)
    for tag, value in zip(tags, data):
        add_run(optional, f'{tag}: {value}')
        if tag != 'Личные качества':
            newline(optional)



def create_docx(name : str, data : dict, photo : str):
    pr_tags = ['Профессия', 'Занятость', 'График работы', 'Желаемая зарплата', 'Телефон', 'Электронная почта']
    pers_tags = ['Дата рождения', 'Пол', 'Гражданство', 'Место проживания', 'Образование', 'Семейное положение']
    profile = [data[i] for i in pr_tags]
    personal = [data[i] for i in pers_tags]
    if data['exp']:
        exp_tags = ['Период работы', 'Должность', 'Организация', 'Должностные обязанности и достижения']
        exp = [data[i] for i in exp_tags]
    else:
        exp = None
        exp_tags = None
    if data['Образование'] in ('высшее', 'вреднее профессиональное'):
        edu_tags = ['Учебное заведение', 'Год окончания', 'Факультет', 'Специальность', 'Форма обучения']
    elif data['Образование'] == 'среднее':
        edu_tags = ['Учебное заведение', 'Год окончания', 'Форма обучения']
    else:
        edu_tags = None
    edu = [data[i] for i in edu_tags] if edu_tags else None
    optional_tags = ['Иностранные языки', 'Наличие водительских прав', 'Служба в армии', 'Увлечения', 'Личные качества']
    optional = [data[i] for i in optional_tags]
    profile.insert(0, data['name'])
    document = Document('sample.docx')
    complete_profile(document, pr_tags, profile, photo)
    complete_personal(document, pers_tags, personal)
    complete_education(document, edu_tags, edu)
    complete_exp(document, exp_tags, exp, data['exp'])
    complete_optional(document, optional_tags, optional)
    document.save(name)


